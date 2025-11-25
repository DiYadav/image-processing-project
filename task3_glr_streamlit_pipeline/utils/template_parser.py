from docx import Document
import re

def extract_placeholders(doc_path):
    """Return list of placeholder keys found like {{key}}"""
    doc = Document(doc_path)
    results = []
    for p in doc.paragraphs:
        results += re.findall(r"\{\{(.*?)\}\}", p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                results += re.findall(r"\{\{(.*?)\}\}", cell.text)
    return list(dict.fromkeys([r.strip() for r in results if r.strip()]))

def extract_labels(doc_path):
    """
    Extract human labels from docx (lines or table cells that look like 'Label:' or 'Label -')
    Returns list of dicts: [{"label":"Client Name","key":"client_name"},...]
    """
    doc = Document(doc_path)
    labels = []
    def process_text(text):
        text = text.strip()
        if not text:
            return None
        # skip if already placeholder present
        if re.search(r"\{\{.*\}\}", text):
            return None
        # heuristic: line contains colon
        if ":" in text:
            left = text.split(":",1)[0].strip()
            if left and len(left) <= 80:
                key = re.sub(r'[^0-9a-zA-Z]+', '_', left).strip('_').lower()
                return {"label": left, "key": key}
        return None

    for p in doc.paragraphs:
        item = process_text(p.text)
        if item:
            labels.append(item)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                item = process_text(cell.text)
                if item:
                    labels.append(item)
    # deduplicate preserving first occurrence
    seen = set(); unique=[]
    for it in labels:
        if it["key"] not in seen:
            seen.add(it["key"])
            unique.append(it)
    return unique
