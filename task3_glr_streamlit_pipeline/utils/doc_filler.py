from docx import Document
import re

def fill_placeholders(template_path, output_path, kv):
    """
    Replace {{key}} placeholders with values in paragraphs and tables.
    """
    doc = Document(template_path)
    # paragraphs
    for p in doc.paragraphs:
        text = p.text
        for k, v in kv.items():
            placeholder = "{{" + k + "}}"
            if placeholder in text:
                text = text.replace(placeholder, v)
        # replace runs
        for run in list(p.runs):
            p._element.remove(run._element)
        p.add_run(text)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    for k, v in kv.items():
                        placeholder = "{{" + k + "}}"
                        if placeholder in text:
                            text = text.replace(placeholder, v)
                    for run in list(p.runs):
                        p._element.remove(run._element)
                    p.add_run(text)

    doc.save(output_path)
    return output_path

def replace_label_in_paragraph(p, label, value):
    # find label followed by optional colon and any trailing text; replace with "Label: value"
    pattern = re.compile(re.escape(label) + r"\s*:?\s*(.*)$", re.IGNORECASE)
    if pattern.search(p.text):
        new_text = pattern.sub(f"{label}: {value}", p.text)
        # replace runs
        for run in list(p.runs):
            p._element.remove(run._element)
        p.add_run(new_text)
        return True
    return False

def fill_labels(template_path, output_path, mappings, kv):
    """
    mappings: list of {"label": "...", "key":"..."}
    kv: {key: value}
    """
    doc = Document(template_path)
    for p in doc.paragraphs:
        for m in mappings:
            key = m["key"]
            label = m["label"]
            if key in kv and kv[key]:
                replace_label_in_paragraph(p, label, kv[key])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in mappings:
                        key = m["key"]
                        label = m["label"]
                        if key in kv and kv[key]:
                            replace_label_in_paragraph(p, label, kv[key])

    doc.save(output_path)
    return output_path

def fill_template(template_path, output_path, kv, mappings=None):
    """
    If mappings provided -> use label-fill mode, else placeholder mode.
    """
    if mappings:
        return fill_labels(template_path, output_path, mappings, kv)
    else:
        return fill_placeholders(template_path, output_path, kv)
